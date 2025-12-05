"""
Document Loader Service
Loads and extracts text from resume documents (.docx format)
"""

import os
from docx import Document
from pathlib import Path


class DocumentLoader:
    """Loads and caches resume content from .docx files"""
    
    _resume_content: str = None
    
    @classmethod
    def load_resume(cls, file_path: str = None) -> str:
        """
        Load resume content from a .docx file
        
        Args:
            file_path: Optional path to resume file. 
                       Defaults to RESUME_PATH env var or data/resume.docx
        
        Returns:
            String containing the full resume text
        """
        if cls._resume_content is not None:
            return cls._resume_content
        
        # Determine file path
        if file_path is None:
            file_path = os.getenv("RESUME_PATH", "data/resume.docx")
        
        # Make path absolute if relative
        if not os.path.isabs(file_path):
            base_dir = Path(__file__).parent.parent
            file_path = base_dir / file_path
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        # Load .docx file
        if file_path.suffix.lower() in ['.docx', '.doc']:
            cls._resume_content = cls._load_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            cls._resume_content = cls._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        return cls._resume_content
    
    @staticmethod
    def _load_docx(file_path: Path) -> str:
        """Extract text from a .docx file"""
        doc = Document(file_path)
        
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    
    @staticmethod
    def _load_txt(file_path: Path) -> str:
        """Load text from a .txt file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @classmethod
    def reload_resume(cls, file_path: str = None) -> str:
        """Force reload of resume content"""
        cls._resume_content = None
        return cls.load_resume(file_path)
    
    @classmethod
    def get_resume_summary(cls) -> dict:
        """Get resume content with metadata"""
        content = cls.load_resume()
        return {
            "content": content,
            "character_count": len(content),
            "word_count": len(content.split())
        }
